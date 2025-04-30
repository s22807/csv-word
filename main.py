import docx
import pandas as pd
import config
import os


#template = docx.Document("wzór karty przekazowej zażalenie.docx")
#paragrapghs = [f"{x.text}" for x in template.paragraphs]
#for p in paragrapghs:
#    p.replace()
#    print(p)


df = pd.read_csv(config.plik_csv, sep=';')
# df.info()
df[df.columns] = df[df.columns].astype(str)
# df.info()



for index,row in df.iterrows():
    template = docx.Document(config.word_template)
    
    for p in template.paragraphs:
        style = p.style
        words_to_replace = [word for word in p.text.split() if (word.endswith('}') or word.endswith('};')) and word.startswith('{')]
        for word in words_to_replace:
            word_no_brackets = "".join(filter(lambda c: c not in ['{','}',';'], word))
            if df.columns.__contains__(word_no_brackets):
                p.text = p.text.replace(word, row[word_no_brackets])
                
        p.style = style
        #print(p.text)

    if config.absolute_dir:
        path = config.absolute_dir
    else:
        path = "./"
    if config.dirs:
        for dir in config.dirs:
            if df.columns.__contains__(dir):
                dir = row[dir]
            path = f"{path}{dir}/"
            if not os.path.isdir(path):
                os.mkdir(path)

    template.save(f'{path}{index+1} uzupelniony.docx')
    



